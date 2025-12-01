# DEPENDENCIES
import docx
import hashlib
from typing import List
from pathlib import Path
from typing import Optional
from docx.table import Table
from datetime import datetime
from docx.document import Document
from config.models import DocumentType
from docx.text.paragraph import Paragraph
from utils.text_cleaner import TextCleaner
from config.models import DocumentMetadata
from config.logging_config import get_logger
from utils.error_handler import handle_errors
from utils.error_handler import DOCXParseError


# Setup Logging
logger = get_logger(__name__)


class DOCXParser:
    """
    Comprehensive DOCX parsing with structure preservation: Handles paragraphs, tables, headers, and footers
    """
    def __init__(self):
        self.logger = logger
    

    @handle_errors(error_type = DOCXParseError, log_error = True, reraise = True)
    def parse(self, file_path: Path, extract_metadata: bool = True, clean_text: bool = True, include_tables: bool = True, include_headers_footers: bool = False) -> tuple[str, Optional[DocumentMetadata]]:
        """
        Parse DOCX and extract text and metadata
        
        Arguments:
        ----------
            file_path               { Path } : Path to DOCX file
            
            extract_metadata        { bool } : Extract document metadata
            
            clean_text              { bool } : Clean extracted text
            
            include_tables          { bool } : Include table content
            
            include_headers_footers { bool } : Include headers and footers
        
        Returns:
        --------
                      { tuple }              : Tuple of (extracted_text, metadata)
        
        Raises:
        -------
            DOCXParseError                   : If parsing fails
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise DOCXParseError(str(file_path), original_error = FileNotFoundError(f"DOCX file not found: {file_path}"))
        
        self.logger.info(f"Parsing DOCX: {file_path}")
        
        try:
            # Open document
            doc            = docx.Document(file_path)
            
            # Extract text content
            text_parts     = list()
            
            # Extract paragraphs
            paragraph_text = self._extract_paragraphs(doc = doc)

            text_parts.append(paragraph_text)
            
            # Extract tables
            if include_tables:
                table_text = self._extract_tables(doc)
                
                if table_text:
                    text_parts.append("\n[TABLES]\n" + table_text)
            
            # Extract headers and footers
            if include_headers_footers:
                header_footer_text = self._extract_headers_footers(doc)
                
                if header_footer_text:
                    text_parts.append("\n[HEADERS/FOOTERS]\n" + header_footer_text)
            
            # Combine all text
            text_content = "\n".join(text_parts)
            
            # Extract metadata
            metadata     = None
            
            if extract_metadata:
                metadata = self._extract_metadata(doc, file_path)
            
            # Clean text
            if clean_text:
                text_content = TextCleaner.clean(text_content,
                                                 remove_html          = False, 
                                                 normalize_whitespace = True,
                                                 preserve_structure   = True,
                                                )
            
            self.logger.info(f"Successfully parsed DOCX: {len(text_content)} characters, {len(doc.paragraphs)} paragraphs")
            
            return text_content, metadata
        
        except Exception as e:
            self.logger.error(f"Failed to parse DOCX {file_path}: {str(e)}")

            raise DOCXParseError(str(file_path), original_error = e)
    

    def _extract_paragraphs(self, doc: Document) -> str:
        """
        Extract text from paragraphs, preserving structure
        
        Arguments:
        ----------
            doc { Document } : Document object
        
        Returns:
        --------
              { str }        : Combined paragraph text
        """
        text_parts = list()
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect headings
            if para.style.name.startswith('Heading'):
                heading_level = para.style.name.replace('Heading', '').strip()
                text_parts.append(f"\n[HEADING {heading_level}] {text}\n")
            
            else:
                text_parts.append(text)
        
        return "\n".join(text_parts)
    

    def _extract_tables(self, doc: Document) -> str:
        """
        Extract text from tables
        
        Arguments:
        ----------
            doc { Document } : Document object
        
        Returns:
        --------
               { str }       : Combined table text
        """
        if not doc.tables:
            return ""
        
        table_parts = list()
        
        for table_idx, table in enumerate(doc.tables):
            table_text = self._parse_table(table)
            
            if table_text:
                table_parts.append(f"\n[TABLE {table_idx + 1}]\n{table_text}")
        
        return "\n".join(table_parts)
    

    def _parse_table(self, table: Table) -> str:
        """
        Parse a single table into text
        
        Arguments:
        ----------
            table { Table } : Table object
        
        Returns:
        --------
              { str }       : Table text
        """
        rows_text = list()
        
        for row in table.rows:
            cells_text = list()
            for cell in row.cells:
                cell_text = cell.text.strip()

                cells_text.append(cell_text)
            
            # Join cells with pipe separator for readability
            rows_text.append(" | ".join(cells_text))
        
        return "\n".join(rows_text)
    

    def _extract_headers_footers(self, doc: Document) -> str:
        """
        Extract headers and footers
        
        Arguments:
        ----------
            doc { Document } : Document object
        
        Returns:
        --------
               { str }       : Headers and footers text
        """
        parts = list()
        
        # Extract from each section
        for section in doc.sections:
            # Header
            if section.header:
                header_text = self._extract_paragraphs_from_element(element = section.header)
                
                if header_text:
                    parts.append(f"[HEADER]\n{header_text}")
            
            # Footer
            if section.footer:
                footer_text = self._extract_paragraphs_from_element(element = section.footer)
                
                if footer_text:
                    parts.append(f"[FOOTER]\n{footer_text}")
        
        return "\n".join(parts)
    

    @staticmethod
    def _extract_paragraphs_from_element(element) -> str:
        """
        Extract paragraphs from header/footer element
        """
        parts = list()

        for para in element.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        return "\n".join(parts)
    

    def _extract_metadata(self, doc: Document, file_path: Path) -> DocumentMetadata:
        """
        Extract metadata from DOCX
        
        Arguments:
        ----------
            doc       { Document} : Document object

            file_path   { Path }  : Path to DOCX file
        
        Returns:
        --------
            { DocumentMetadata }  : DocumentMetadata object
        """
        # Get core properties
        core_props      = doc.core_properties
        
        # Extract fields
        title           = core_props.title or file_path.stem
        author          = core_props.author
        created_date    = core_props.created
        modified_date   = core_props.modified
        
        # Get file size
        file_size       = file_path.stat().st_size
        
        # Generate document ID
        doc_hash        = hashlib.md5(str(file_path).encode()).hexdigest()
        doc_id          = f"doc_{int(datetime.now().timestamp())}_{doc_hash}"
        
        # Count paragraphs and estimate pages
        num_paragraphs  = len(doc.paragraphs)
        
        # Rough estimate: 500 words per page, 5-10 words per paragraph
        estimated_pages = max(1, num_paragraphs // 50)
        
        # Create metadata object
        metadata        = DocumentMetadata(document_id     = doc_id,
                                           filename        = file_path.name,
                                           file_path       = file_path,
                                           document_type   = DocumentType.DOCX,
                                           title           = title,
                                           author          = author,
                                           created_date    = created_date,
                                           modified_date   = modified_date,
                                           file_size_bytes = file_size,
                                           num_pages       = estimated_pages,
                                           extra           = {"num_paragraphs" : num_paragraphs,
                                                              "num_tables"     : len(doc.tables),
                                                              "num_sections"   : len(doc.sections),
                                                              "category"       : core_props.category,
                                                              "comments"       : core_props.comments,
                                                              "keywords"       : core_props.keywords,
                                                              "subject"        : core_props.subject,
                                                             }
                                          )
        
        return metadata
    

    def get_paragraph_count(self, file_path: Path) -> int:
        """
        Get number of paragraphs in document
        
        Arguments:
        ----------
            file_path { Path } : Path to DOCX file
        
        Returns:
        --------
                { int }        : Number of paragraphs
        """
        try:
            doc = docx.Document(file_path)
            return len(doc.paragraphs)

        except Exception as e:
            self.logger.error(f"Failed to get paragraph count: {repr(e)}")
            raise DOCXParseError(str(file_path), original_error = e)
    

    def extract_section(self, file_path: Path, section_index: int, clean_text: bool = True) -> str:
        """
        Extract text from a specific section
        
        Arguments:
        ----------
            file_path     { Path } : Path to DOCX file
            
            section_index { int }  : Section index (0-indexed)
            
            clean_text    { bool } : Clean extracted text
        
        Returns:
        --------
                    { str }        : Section text
        """
        try:
            doc = docx.Document(file_path)
            
            if ((section_index < 0) or (section_index >= len(doc.sections))):
                raise ValueError(f"Section index {section_index} out of range (0-{len(doc.sections)-1})")
            
            # Note: Extracting text by section is not straightforward in python-docx
            section = doc.sections[section_index]
            
            # For now, we'll extract the entire document
            text    = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if clean_text:
                text = TextCleaner.clean(text)
            
            return text
        
        except Exception as e:
            self.logger.error(f"Failed to extract section: {repr(e)}")
            raise DOCXParseError(str(file_path), original_error = e)
    

    def extract_heading_sections(self, file_path: Path, clean_text: bool = True) -> dict[str, str]:
        """
        Extract text organized by headings
        
        Arguments:
        ----------
            file_path  { Path } : Path to DOCX file

            clean_text { bool } : Clean extracted text
        
        Returns:
        --------
                { dict }        : Dictionary mapping heading text to content
        """
        try:
            doc             = docx.Document(file_path)
            sections        = dict()
            current_content = list()
            current_heading = "Introduction"
            
            
            for para in doc.paragraphs:
                text = para.text.strip()
                
                if not text:
                    continue
                
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    # Save previous section
                    if current_content:
                        section_text = "\n".join(current_content)
                        
                        if clean_text:
                            section_text = TextCleaner.clean(section_text)
                        
                        sections[current_heading] = section_text
                    
                    # Start new section
                    current_heading = text
                    current_content = list()

                else:
                    current_content.append(text)
            
            # Save last section
            if current_content:
                section_text = "\n".join(current_content)
                
                if clean_text:
                    section_text = TextCleaner.clean(section_text)
                
                sections[current_heading] = section_text
            
            return sections
        
        except Exception as e:
            self.logger.error(f"Failed to extract heading sections: {repr(e)}")
            raise DOCXParseError(str(file_path), original_error = e)
